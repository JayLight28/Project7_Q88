import os
import datetime
import re
import shutil
import socket
import threading
import time
import tkinter as tk
import uuid
import webbrowser
from tkinter import filedialog

import docx
from flask import Flask, render_template, request, redirect, url_for, jsonify, g, abort

from q88 import parser, rules, state as statemod, style as stylemod, locks, config as configmod

APP_DIR = os.path.dirname(os.path.abspath(__file__))
APP_VERSION = "1.2.1"
DEFAULT_WARNING_DAYS = 60
REFERENCE_HINT = "original form"
PORT = 5000
CLIENT_COOKIE = "q88_client_id"
NAME_COOKIE = "q88_name"
FOLDER_COOKIE = "q88_folder"

# "Q88 V6 <ship code> <date>.docx" - only the date changes when we rename on save
FILENAME_RE = re.compile(r"^(Q88 V6 )([^ ]+)( )(.+)(\.docx)$", re.IGNORECASE)

app = Flask(__name__)
# The Q88 questionnaire form has thousands of fields (one per cell), well past
# Werkzeug's default max_form_parts=1000 - without this, /save 413s on any
# full-size document.
app.config["MAX_FORM_PARTS"] = 10000

# in-memory cache of the most recently opened document, keyed by "folder::filename"
# (folder is per-browser now, so two people can look at different folders at once
# without colliding even if a filename happens to repeat), so /save can locate the
# exact python-docx Cell objects parsed by /open. Guarded by _cache_mutex since
# waitress serves requests from multiple worker threads.
_OPEN_CACHE = {}
_cache_mutex = threading.Lock()

# home-page file-list cache (issue count + vessel name/flag), keyed by
# (path, mtime) - parsing a large .docx and classifying every cell costs
# ~0.5-1s; without this, the home page re-did that for every file on every
# load (10 files -> 10+ seconds).
_FILE_SCAN_CACHE = {}
_file_scan_mutex = threading.Lock()


def _key(folder, filename):
    return f"{folder}::{filename}"


def get_client_id():
    """Stable per-browser id. Generated on first request and persisted as a
    cookie so it's available immediately within this same request too."""
    cid = request.cookies.get(CLIENT_COOKIE)
    if not cid:
        cid = uuid.uuid4().hex
        g._new_client_id = cid
    return cid


def get_display_name():
    return request.cookies.get(NAME_COOKIE) or "Anonymous"


def get_current_folder():
    """Each browser can point at its own folder (files stay wherever they are -
    e.g. a network share - only the folder *selection* is per-user)."""
    folder = request.cookies.get(FOLDER_COOKIE)
    if folder and os.path.isdir(folder):
        return folder
    return configmod.get_watch_folder()


@app.context_processor
def _inject_version():
    return {"app_version": APP_VERSION}


@app.after_request
def _ensure_client_cookie(response):
    new_cid = getattr(g, "_new_client_id", None)
    if new_cid:
        response.set_cookie(CLIENT_COOKIE, new_cid, max_age=60 * 60 * 24 * 365)
    return response


def _safe_path(folder, filename):
    """Join folder+filename and refuse to return anything outside folder -
    filename comes straight from the URL (Flask's <path:...> converter
    allows slashes and "..") so this is the one place that containment is
    enforced for every route that opens/saves a vessel file."""
    folder_real = os.path.realpath(folder)
    candidate = os.path.realpath(os.path.join(folder, filename))
    if os.path.commonpath([folder_real, candidate]) != folder_real:
        abort(404)
    return candidate


def list_files(folder):
    files = []
    for name in sorted(os.listdir(folder)):
        lower = name.lower()
        if lower.endswith(".original_backup.docx") or lower.endswith(".original_backup.doc"):
            continue
        mtime = int(os.path.getmtime(os.path.join(folder, name)))
        if lower.endswith(".docx"):
            files.append({"name": name, "supported": True, "mtime": mtime})
        elif lower.endswith(".doc"):
            files.append({"name": name, "supported": False, "mtime": mtime})
    return files


def open_document(filename, folder):
    """Re-reads the .docx from disk unless the cache already holds a parse
    of this exact on-disk version (same mtime) - e.g. right after this
    process itself just wrote the file in add_row/delete_row, where
    re-reading+re-unzipping over a network share is pure overhead."""
    path = _safe_path(folder, filename)
    key = _key(folder, filename)
    try:
        mtime = os.path.getmtime(path)
    except OSError:
        mtime = None

    cached = _OPEN_CACHE.get(key)
    if cached is not None and mtime is not None and cached.get("mtime") == mtime:
        return cached["ext"]

    doc = docx.Document(path)
    ext = parser.extract(doc)
    _OPEN_CACHE[key] = {"path": path, "doc": doc, "ext": ext, "mtime": mtime}
    return ext


def find_reference_file(files):
    for f in files:
        if f["supported"] and REFERENCE_HINT in f["name"].lower():
            return f["name"]
    return None


def _compute_issues(path, warning_days=DEFAULT_WARNING_DAYS, ext=None):
    """Every currently-flagged field in a document, newest/most-severe first.
    Shared by the file-list issue count and the home-page detail panel.

    A multi-column table row (e.g. one row, five columns) collapses to at
    most one issue - if every cell in the row is filled (or N/A-checked),
    the row has no issue; otherwise one issue represents the whole row,
    at its most severe cell's state.

    Pass a pre-parsed `ext` (e.g. from `_quick_scan`) to avoid re-parsing the
    same .docx twice in one request."""
    if ext is None:
        ext = parser.extract(docx.Document(path))
    st = statemod.load_state(path)
    today = datetime.date.today()
    state_order = {"EXPIRED": 0, "WARNING": 1, "MISSING": 2}

    def cell_state(fid, label, col):
        cell = ext.cell_map.get(fid)
        if cell is None or st["na_flags"].get(fid):
            return None, ""
        state, _ = rules.classify(label, col, cell.text.strip(), warning_days=warning_days, today=today)
        return state, cell.text.strip()

    issues = []
    for r in ext.display_rows:
        if r["type"] == "field":
            state, text = cell_state(r["id"], r["label"], r["column_header"])
            if state in rules.HIGHLIGHTABLE:
                full_label = f"{r['label']} ({r['column_header']})" if r["column_header"] else r["label"]
                issues.append({"id": r["id"], "item_code": r["item_code"], "label": full_label, "text": text, "state": state})
        elif r["type"] == "table":
            for tr in r["rows"]:
                item_code = tr.get("row_item_code") or r["item_code"]
                row_flagged = []
                for c in tr["cells"]:
                    state, _ = cell_state(c["id"], tr["row_label"], c["column_header"])
                    if state in rules.HIGHLIGHTABLE:
                        row_flagged.append((state, c["id"]))
                if row_flagged:
                    worst_state = min((s for s, _ in row_flagged), key=lambda s: state_order.get(s, 9))
                    issues.append({"id": row_flagged[0][1], "item_code": item_code, "label": tr["row_label"], "text": "", "state": worst_state})

    issues.sort(key=lambda x: state_order.get(x["state"], 9))
    return issues


def _quick_scan(path):
    """One parse per (path, mtime) for the home-page file cards: how many
    fields are flagged, plus the vessel name/IMO (item 1.2) and flag/port of
    registry (item 1.5) so a vessel is identifiable without opening the file.
    Cached so the home page doesn't re-parse every file's full .docx on every
    load - that alone used to take 10+ seconds across a 10-file folder."""
    try:
        mtime = os.path.getmtime(path)
    except OSError:
        return None

    cache_key = (path, mtime)
    with _file_scan_mutex:
        if cache_key in _FILE_SCAN_CACHE:
            return _FILE_SCAN_CACHE[cache_key]

    try:
        ext = parser.extract(docx.Document(path))
        issue_count = len(_compute_issues(path, ext=ext))
        vessel_name = ""
        flag = ""
        for r in ext.display_rows:
            if r["type"] != "field":
                continue
            if r["item_code"] == "1.2":
                vessel_name = ext.cell_map[r["id"]].text.strip()
            elif r["item_code"] == "1.5":
                flag = ext.cell_map[r["id"]].text.strip()
        result = {"issue_count": issue_count, "vessel_name": vessel_name, "flag": flag}
    except Exception:
        return None

    with _file_scan_mutex:
        for stale_key in [k for k in _FILE_SCAN_CACHE if k[0] == path]:
            del _FILE_SCAN_CACHE[stale_key]
        _FILE_SCAN_CACHE[cache_key] = result
    return result


def _native_folder_dialog(initialdir):
    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    try:
        path = filedialog.askdirectory(initialdir=initialdir or None)
    finally:
        root.destroy()
    return path or None


def _native_file_dialog(initialdir):
    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    try:
        path = filedialog.askopenfilename(
            initialdir=initialdir or None,
            filetypes=[("Word documents", "*.docx *.doc"), ("All files", "*.*")],
        )
    finally:
        root.destroy()
    return path or None


@app.route("/")
def index():
    folder = get_current_folder()
    files = list_files(folder)
    for f in files:
        scan = _quick_scan(os.path.join(folder, f["name"])) if f["supported"] else None
        f["issue_count"] = scan["issue_count"] if scan else None
        f["vessel_name"] = scan["vessel_name"] if scan else ""
        f["flag"] = scan["flag"] if scan else ""
    reference = find_reference_file(files)
    fleets = configmod.get_fleets()
    active_fleet = next(
        (name for name, p in fleets.items()
         if os.path.normcase(os.path.normpath(p)) == os.path.normcase(os.path.normpath(folder))),
        None,
    )
    return render_template(
        "index.html", files=files, reference=reference,
        style_applied=request.args.get("style_applied"),
        style_count=request.args.get("style_count"),
        style_files=request.args.get("style_files"),
        display_name=get_display_name(),
        watch_folder=folder,
        fleets=fleets,
        active_fleet=active_fleet,
        folder_saved=request.args.get("folder_saved") == "1",
        imported=request.args.get("imported"),
        renamed=request.args.get("renamed"),
        rename_error=request.args.get("rename_error"),
    )


@app.route("/pick_folder", methods=["POST"])
def pick_folder():
    typed = (request.form.get("folder_path") or "").strip()
    if typed:
        chosen = typed
    else:
        chosen = _native_folder_dialog(get_current_folder())
    if not chosen or not os.path.isdir(chosen):
        return redirect(url_for("index", rename_error="bad_folder"))
    folder = os.path.normpath(chosen)

    make_default = request.form.get("make_default") == "on"
    if make_default:
        configmod.set_watch_folder(folder)
        _write_shortcut(folder)

    resp = redirect(url_for("index", folder_saved=1))
    resp.set_cookie(FOLDER_COOKIE, folder, max_age=60 * 60 * 24 * 365)
    return resp


@app.route("/import_file", methods=["POST"])
def import_file():
    folder = get_current_folder()
    upload = request.files.get("file")
    if upload and upload.filename:
        dest_name = os.path.basename(upload.filename)
        if not dest_name.lower().endswith((".docx", ".doc")):
            return redirect(url_for("index", rename_error="bad_import"))
        dest_path = os.path.join(folder, dest_name)
        if os.path.exists(dest_path):
            base, ext = os.path.splitext(dest_name)
            i = 1
            while os.path.exists(os.path.join(folder, f"{base} ({i}){ext}")):
                i += 1
            dest_name = f"{base} ({i}){ext}"
            dest_path = os.path.join(folder, dest_name)
        upload.save(dest_path)
        return redirect(url_for("index", imported=dest_name))

    chosen = _native_file_dialog(folder)
    if not chosen:
        return redirect(url_for("index"))

    dest_name = os.path.basename(chosen)
    dest_path = os.path.join(folder, dest_name)
    if os.path.abspath(chosen) != os.path.abspath(dest_path):
        if os.path.exists(dest_path):
            base, ext = os.path.splitext(dest_name)
            i = 1
            while os.path.exists(os.path.join(folder, f"{base} ({i}){ext}")):
                i += 1
            dest_name = f"{base} ({i}){ext}"
            dest_path = os.path.join(folder, dest_name)
        shutil.copy2(chosen, dest_path)

    return redirect(url_for("index", imported=dest_name))


@app.route("/rename_file/<path:filename>", methods=["POST"])
def rename_file(filename):
    folder = get_current_folder()
    lock_key = _key(folder, filename)
    client_id = get_client_id()
    if not locks.is_owner(lock_key, client_id):
        return redirect(url_for("index", rename_error="locked"))

    new_name = os.path.basename((request.form.get("new_name") or "").strip())
    if not new_name:
        return redirect(url_for("index", rename_error="empty"))
    if not new_name.lower().endswith(".docx"):
        new_name += ".docx"

    path = _safe_path(folder, filename)
    new_path = os.path.join(folder, new_name)
    if os.path.exists(new_path):
        return redirect(url_for("index", rename_error="exists"))

    _, ok = _manual_rename(path, new_name, get_display_name())
    if ok:
        _OPEN_CACHE.pop(lock_key, None)
        with _file_scan_mutex:
            for stale_key in [k for k in _FILE_SCAN_CACHE if k[0] == path]:
                del _FILE_SCAN_CACHE[stale_key]

    return redirect(url_for("index", renamed=new_name if ok else None))


@app.route("/set_name", methods=["POST"])
def set_name():
    resp = redirect(request.form.get("next") or url_for("index"))
    name = (request.form.get("display_name") or "").strip()[:30] or "Anonymous"
    resp.set_cookie(NAME_COOKIE, name, max_age=60 * 60 * 24 * 365)
    return resp


@app.route("/open/<path:filename>")
def open_file(filename):
    if not filename.lower().endswith(".docx"):
        return redirect(url_for("index"))

    folder = get_current_folder()
    lock_key = _key(folder, filename)
    client_id = get_client_id()
    name = get_display_name()
    got_lock, holder = locks.acquire(lock_key, client_id, name)
    read_only = not got_lock

    ext = open_document(filename, folder)
    path = _safe_path(folder, filename)
    st = statemod.load_state(path)
    warning_days = int(request.args.get("warning_days", DEFAULT_WARNING_DAYS))
    today = datetime.date.today()

    issues = []
    recently_changed = set(st.get("last_changed_ids", []))

    state_order = {"EXPIRED": 0, "WARNING": 1, "MISSING": 2}

    def classify_cell(cell_rec, label, column_header, item_code, collect=True):
        computed_state, _ = rules.classify(
            label, column_header, cell_rec["text"],
            warning_days=warning_days, today=today,
        )
        na_checked = bool(st["na_flags"].get(cell_rec["id"]))
        display_state = "OK" if (na_checked and computed_state in rules.HIGHLIGHTABLE) else computed_state
        cell_rec["display_state"] = display_state
        cell_rec["na_checked"] = na_checked
        cell_rec["show_na_checkbox"] = computed_state in rules.HIGHLIGHTABLE
        cell_rec["recently_changed"] = cell_rec["id"] in recently_changed
        if collect and display_state in rules.HIGHLIGHTABLE:
            full_label = f"{label} ({column_header})" if column_header else label
            issues.append({
                "id": cell_rec["id"], "item_code": item_code,
                "label": full_label, "text": cell_rec["text"], "state": display_state,
            })
        return display_state

    rows = []
    sections = []  # sidebar nav: [{id, text, kind: 'heading'|'subheading'}]
    section_id = "sec-0"
    section_counter = 0
    for r in ext.display_rows:
        if r["type"] == "heading":
            section_counter += 1
            section_id = f"sec-{section_counter}"
            sections.append({"id": section_id, "text": r["text"], "kind": "heading"})
        elif r["type"] == "subheading":
            sections.append({"id": f"{section_id}-{len(sections)}", "text": r["text"], "kind": "subheading", "parent": section_id})

        if r["type"] == "field":
            row = dict(r)
            row["section_id"] = section_id
            classify_cell(row, row["label"], row["column_header"], row["item_code"])
            rows.append(row)
        elif r["type"] == "table":
            row = dict(r)
            row["section_id"] = section_id
            row["rows"] = []
            for tr in r["rows"]:
                trc = dict(tr)
                trc["cells"] = [dict(c) for c in tr["cells"]]
                item_code = tr.get("row_item_code") or r["item_code"]
                row_flagged = []
                for c in trc["cells"]:
                    cell_state = classify_cell(c, tr["row_label"], c["column_header"], item_code, collect=False)
                    if cell_state in rules.HIGHLIGHTABLE:
                        row_flagged.append((cell_state, c["id"]))
                if row_flagged:
                    worst_state = min((s for s, _ in row_flagged), key=lambda s: state_order.get(s, 9))
                    issues.append({
                        "id": row_flagged[0][1], "item_code": item_code,
                        "label": tr["row_label"], "text": "", "state": worst_state,
                    })
                row["rows"].append(trc)
            rows.append(row)
        else:
            row = dict(r)
            row["section_id"] = section_id
            if r["type"] == "subheading":
                row["anchor_id"] = sections[-1]["id"]
            rows.append(row)

    issues.sort(key=lambda x: state_order.get(x["state"], 9))

    return render_template(
        "document.html",
        filename=filename,
        rows=rows,
        sections=sections,
        issues=issues,
        warning_days=warning_days,
        saved=request.args.get("saved") == "1",
        read_only=read_only,
        lock_holder_name=(holder or {}).get("name"),
        conflict=request.args.get("conflict") == "1",
        saved_as=request.args.get("saved_as"),
        save_as_error=request.args.get("save_as_error"),
    )


@app.route("/heartbeat/<path:filename>", methods=["POST"])
def heartbeat(filename):
    lock_key = _key(get_current_folder(), filename)
    client_id = get_client_id()
    ok, holder = locks.acquire(lock_key, client_id, get_display_name())
    return jsonify({"ok": ok, "holder": (holder or {}).get("name")})


@app.route("/lock_status/<path:filename>")
def lock_status(filename):
    lock_key = _key(get_current_folder(), filename)
    holder = locks.status(lock_key)
    return jsonify({"locked": bool(holder), "holder": (holder or {}).get("name")})


@app.route("/release/<path:filename>", methods=["POST"])
def release_lock(filename):
    lock_key = _key(get_current_folder(), filename)
    locks.release(lock_key, get_client_id())
    return jsonify({"ok": True})


@app.route("/panel/<path:filename>")
def file_panel(filename):
    folder = get_current_folder()
    path = _safe_path(folder, filename)
    issues = _compute_issues(path)
    st = statemod.load_state(path)
    raw = st.get("history", [])
    entries = [dict(e, index=i) for i, e in enumerate(raw)]
    entries.reverse()
    return render_template(
        "_panel.html",
        filename=filename,
        issues=issues,
        history=entries,
        has_backup=os.path.exists(statemod.backup_path(path)),
    )


@app.route("/field_edit/<path:filename>/<field_id>", methods=["POST"])
def field_edit(filename, field_id):
    folder = get_current_folder()
    path = _safe_path(folder, filename)
    cache = _get_or_load_cache(filename, folder)
    doc, ext = cache["doc"], cache["ext"]

    cell = ext.cell_map.get(field_id)
    if cell is None:
        return jsonify({"ok": False, "error": "field not found"}), 404

    st = statemod.load_state(path)
    new_text = (request.form.get("text") or "").strip()
    old_text = cell.text.strip()
    if new_text != old_text:
        label = _field_labels(ext).get(field_id, "")
        statemod.ensure_backup(path)
        parser.set_cell_text(cell, new_text)
        statemod.record_edit(st, field_id, label, old_text, new_text, by=get_display_name())
        doc.save(path)

    st["na_flags"][field_id] = request.form.get("na") == "on"
    statemod.save_state(path, st)
    _OPEN_CACHE.pop(_key(folder, filename), None)

    return jsonify({"ok": True})


@app.route("/history/<path:filename>")
def history(filename):
    path = os.path.join(get_current_folder(), filename)
    st = statemod.load_state(path)
    raw = st.get("history", [])
    entries = [dict(e, index=i) for i, e in enumerate(raw)]
    entries.reverse()
    has_backup = os.path.exists(statemod.backup_path(path))
    return render_template(
        "history.html", filename=filename, entries=entries, has_backup=has_backup,
        restored=request.args.get("restored") == "1",
        reverted=request.args.get("reverted") == "1",
    )


def _wants_ajax():
    return request.args.get("ajax") == "1" or request.form.get("ajax") == "1"


@app.route("/history/<path:filename>/revert/<int:index>", methods=["POST"])
def revert_entry(filename, index):
    folder = get_current_folder()
    path = _safe_path(folder, filename)
    st = statemod.load_state(path)
    raw = st.get("history", [])
    if index < 0 or index >= len(raw):
        if _wants_ajax():
            return jsonify({"ok": False}), 404
        return redirect(url_for("history", filename=filename))

    entry = raw[index]
    field_id = entry.get("field_id")
    if field_id and field_id not in ("__rename__", "__restore__"):
        doc = docx.Document(path)
        ext = parser.extract(doc)
        cell = ext.cell_map.get(field_id)
        if cell is not None:
            current_text = cell.text.strip()
            old_text = entry.get("old", "")
            if current_text != old_text:
                statemod.ensure_backup(path)
                parser.set_cell_text(cell, old_text)
                statemod.record_edit(
                    st, field_id, entry.get("label", ""), current_text, old_text,
                    by=f"{get_display_name()} (revert)",
                )
                doc.save(path)
                statemod.save_state(path, st)
                _OPEN_CACHE.pop(_key(folder, filename), None)

    if _wants_ajax():
        return jsonify({"ok": True})
    return redirect(url_for("history", filename=filename, reverted=1))


@app.route("/restore_original/<path:filename>", methods=["POST"])
def restore_original(filename):
    folder = get_current_folder()
    path = _safe_path(folder, filename)
    bpath = statemod.backup_path(path)
    if os.path.exists(bpath):
        shutil.copy2(bpath, path)
        st = statemod.load_state(path)
        statemod.record_edit(
            st, "__restore__", "Whole document", "(edited)", "(restored to original)",
            by=get_display_name(),
        )
        statemod.save_state(path, st)
        _OPEN_CACHE.pop(_key(folder, filename), None)
    if _wants_ajax():
        return jsonify({"ok": True})
    return redirect(url_for("history", filename=filename, restored=1))


def _field_labels(ext):
    labels = {}
    for r in ext.display_rows:
        if r["type"] == "field":
            labels[r["id"]] = r["label"]
        elif r["type"] == "table":
            for tr in r["rows"]:
                for c in tr["cells"]:
                    labels[c["id"]] = f"{r['label']} - {tr['row_label']} ({c['column_header']})"
    return labels


def _get_or_load_cache(filename, folder):
    key = _key(folder, filename)
    with _cache_mutex:
        cache = _OPEN_CACHE.get(key)
        if not cache:
            open_document(filename, folder)
            cache = _OPEN_CACHE[key]
        return cache


def _apply_form_edits(ext, path, form, by):
    """Write every f_<id>/na_<id> field present in the submitted form back into
    the in-memory document. Returns True if any cell text actually changed."""
    st = statemod.load_state(path)
    labels_by_id = _field_labels(ext)

    changed = False
    changed_ids = []
    for fid, cell in ext.cell_map.items():
        new_text = form.get(f"f_{fid}")
        if new_text is None:
            continue
        old_text = cell.text.strip()
        new_text = new_text.strip()
        if new_text != old_text:
            if not changed:
                statemod.ensure_backup(path)
            parser.set_cell_text(cell, new_text)
            statemod.record_edit(st, fid, labels_by_id.get(fid, ""), old_text, new_text, by=by)
            changed = True
            changed_ids.append(fid)

    if changed_ids:
        st["last_changed_ids"] = changed_ids

    for fid in labels_by_id:
        st["na_flags"][fid] = form.get(f"na_{fid}") == "on"

    statemod.save_state(path, st)
    return changed


def _rename_for_date(path, new_date_text, by):
    dirname, base = os.path.split(path)
    m = FILENAME_RE.match(base)
    if not m:
        return path
    prefix, code, sep, _old_date, ext = m.groups()
    new_base = f"{prefix}{code}{sep}{new_date_text}{ext}"
    if new_base == base:
        return path
    new_path = os.path.join(dirname, new_base)
    if os.path.exists(new_path):
        return path

    st = statemod.load_state(path)
    statemod.record_edit(st, "__rename__", "Filename", base, new_base, by=by)
    statemod.save_state(path, st)

    os.rename(path, new_path)
    old_state = statemod.state_path(path)
    if os.path.exists(old_state):
        os.rename(old_state, statemod.state_path(new_path))
    return new_path


def _manual_rename(path, new_name, by):
    """User-triggered rename from the home page (any new name, not just a
    date-field update) - also moves the state/backup sidecar files so
    history and 'restore original' keep working under the new name."""
    dirname, base = os.path.split(path)
    new_path = os.path.join(dirname, new_name)
    if new_path == path or os.path.exists(new_path):
        return path, False

    st = statemod.load_state(path)
    statemod.record_edit(st, "__rename__", "Filename", base, new_name, by=by)
    statemod.save_state(path, st)

    os.rename(path, new_path)
    old_state = statemod.state_path(path)
    if os.path.exists(old_state):
        os.rename(old_state, statemod.state_path(new_path))
    old_backup = statemod.backup_path(path)
    if os.path.exists(old_backup):
        os.rename(old_backup, statemod.backup_path(new_path))
    return new_path, True


def _archive_previous_version(path):
    """Copy the pre-edit file into an Obsolete/ subfolder next to it before a
    save overwrites it in place, so prior revisions aren't lost."""
    if not os.path.exists(path):
        return
    dirname, base = os.path.split(path)
    obsolete_dir = os.path.join(dirname, "Obsolete")
    os.makedirs(obsolete_dir, exist_ok=True)
    dest = os.path.join(obsolete_dir, base)
    if os.path.exists(dest):
        stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        name, ext = os.path.splitext(base)
        dest = os.path.join(obsolete_dir, f"{name} ({stamp}){ext}")
    shutil.copy2(path, dest)


def _maybe_rename_by_date_field(ext, path, by):
    """Filenames follow 'Q88 V6 <code> <date>.docx' - keep the date in sync
    with the document's own '1.1 Date updated' field after a save."""
    date_field = next(
        (r for r in ext.display_rows if r.get("type") == "field" and r.get("item_code") == "1.1"),
        None,
    )
    if not date_field:
        return path
    cell = ext.cell_map.get(date_field["id"])
    if cell is None:
        return path
    parsed = rules.try_parse_date(cell.text.strip())
    if not parsed:
        return path
    return _rename_for_date(path, parsed.strftime("%d %B %Y"), by)


@app.route("/save/<path:filename>", methods=["POST"])
def save_file(filename):
    folder = get_current_folder()
    lock_key = _key(folder, filename)
    client_id = get_client_id()
    warning_days = int(request.form.get("warning_days", DEFAULT_WARNING_DAYS))
    if not locks.is_owner(lock_key, client_id):
        return redirect(url_for("open_file", filename=filename, warning_days=warning_days, conflict=1))

    cache = _get_or_load_cache(filename, folder)
    doc, ext, path = cache["doc"], cache["ext"], cache["path"]

    changed = _apply_form_edits(ext, path, request.form, get_display_name())
    if changed:
        _archive_previous_version(path)
        doc.save(path)

    new_path = _maybe_rename_by_date_field(ext, path, get_display_name())
    new_filename = filename
    if new_path != path:
        new_filename = os.path.basename(new_path)
        _OPEN_CACHE.pop(lock_key, None)

    locks.release(lock_key, client_id)

    return redirect(url_for("open_file", filename=new_filename, warning_days=warning_days, saved=1))


@app.route("/save_as/<path:filename>", methods=["POST"])
def save_as(filename):
    """Save the currently-edited (not-yet-saved) form content into a brand
    new .docx file, leaving the original file on disk untouched."""
    folder = get_current_folder()
    lock_key = _key(folder, filename)
    client_id = get_client_id()
    warning_days = int(request.form.get("warning_days", DEFAULT_WARNING_DAYS))
    if not locks.is_owner(lock_key, client_id):
        return redirect(url_for("open_file", filename=filename, warning_days=warning_days, conflict=1))

    new_name = os.path.basename((request.form.get("new_filename") or "").strip())
    if not new_name:
        return redirect(url_for("open_file", filename=filename, warning_days=warning_days, save_as_error="empty"))
    if not new_name.lower().endswith(".docx"):
        new_name += ".docx"

    path = _safe_path(folder, filename)
    new_path = os.path.join(folder, new_name)
    if os.path.exists(new_path):
        return redirect(url_for("open_file", filename=filename, warning_days=warning_days, save_as_error="exists"))

    fresh_doc = docx.Document(path)
    fresh_doc.save(new_path)
    fresh_ext = parser.extract(fresh_doc)
    if _apply_form_edits(fresh_ext, new_path, request.form, get_display_name()):
        fresh_doc.save(new_path)

    return redirect(url_for("open_file", filename=filename, warning_days=warning_days, saved_as=new_name))


@app.route("/add_row/<path:filename>/<table_key>", methods=["POST"])
def add_row(filename, table_key):
    folder = get_current_folder()
    lock_key = _key(folder, filename)
    client_id = get_client_id()
    warning_days = int(request.form.get("warning_days", DEFAULT_WARNING_DAYS))
    if not locks.is_owner(lock_key, client_id):
        return redirect(url_for("open_file", filename=filename, warning_days=warning_days, conflict=1))

    cache = _get_or_load_cache(filename, folder)
    doc, ext, path = cache["doc"], cache["ext"], cache["path"]

    _apply_form_edits(ext, path, request.form, get_display_name())
    nested_table = ext.table_objs.get(table_key)
    if nested_table is not None:
        statemod.ensure_backup(path)
        parser.add_table_row(nested_table)
        doc.save(path)
        # Re-parse the doc object already in memory instead of popping the
        # cache and forcing the next /open to re-read+re-unzip the file from
        # disk (slow over a network share) - this was the main cost of
        # add/delete row on large documents.
        _OPEN_CACHE[lock_key] = {
            "path": path, "doc": doc, "ext": parser.extract(doc),
            "mtime": os.path.getmtime(path),
        }

    return redirect(url_for("open_file", filename=filename, warning_days=warning_days))


@app.route("/delete_row/<path:filename>/<table_key>/<int:row_index>", methods=["POST"])
def delete_row(filename, table_key, row_index):
    folder = get_current_folder()
    lock_key = _key(folder, filename)
    client_id = get_client_id()
    warning_days = int(request.form.get("warning_days", DEFAULT_WARNING_DAYS))
    if not locks.is_owner(lock_key, client_id):
        return redirect(url_for("open_file", filename=filename, warning_days=warning_days, conflict=1))

    cache = _get_or_load_cache(filename, folder)
    doc, ext, path = cache["doc"], cache["ext"], cache["path"]

    _apply_form_edits(ext, path, request.form, get_display_name())
    nested_table = ext.table_objs.get(table_key)
    if nested_table is not None:
        statemod.ensure_backup(path)
        parser.delete_table_row(nested_table, row_index)
        doc.save(path)
        _OPEN_CACHE[lock_key] = {
            "path": path, "doc": doc, "ext": parser.extract(doc),
            "mtime": os.path.getmtime(path),
        }

    return redirect(url_for("open_file", filename=filename, warning_days=warning_days))


def _apply_style_to_file(reference_path, target_filename, folder):
    ref_doc = docx.Document(reference_path)
    style_map = stylemod.extract_style_map(ref_doc)

    target_path = os.path.join(folder, target_filename)
    target_doc = docx.Document(target_path)
    applied = stylemod.apply_style_map(target_doc, style_map)
    if applied:
        statemod.ensure_backup(target_path)
        target_doc.save(target_path)
        _OPEN_CACHE.pop(_key(folder, target_filename), None)
    return applied


@app.route("/apply_style_all", methods=["POST"])
def apply_style_all():
    folder = get_current_folder()
    files = list_files(folder)
    reference = find_reference_file(files)
    total = 0
    file_count = 0
    if reference:
        for f in files:
            if f["supported"] and f["name"] != reference:
                total += _apply_style_to_file(os.path.join(folder, reference), f["name"], folder)
                file_count += 1
    return redirect(url_for("index", style_applied="all files", style_count=total, style_files=file_count))


def _lan_ip():
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        s.connect(("8.8.8.8", 80))
        return s.getsockname()[0]
    except OSError:
        return "127.0.0.1"
    finally:
        s.close()


def _open_browser():
    time.sleep(1.2)
    webbrowser.open(f"http://localhost:{PORT}/")
    # opening the browser steals focus and drops this console window behind
    # it, making the server look like it "went to background" - bring it
    # back to the front so it's obviously still running.
    time.sleep(0.8)
    try:
        import ctypes
        hwnd = ctypes.windll.kernel32.GetConsoleWindow()
        if hwnd:
            ctypes.windll.user32.SetForegroundWindow(hwnd)
    except Exception:
        pass


def _write_shortcut(folder=None):
    """Drop a double-clickable "Open Q88 Check.url" shortcut into the given
    (or default) folder pointing at this PC's current LAN address, so
    colleagues who can already see that folder (e.g. a network share) can
    open the app without typing an IP - and it stays correct if the IP
    changes on the next start."""
    try:
        folder = folder or configmod.get_watch_folder()
        url = f"http://{_lan_ip()}:{PORT}/"
        content = f"[InternetShortcut]\nURL={url}\n"
        with open(os.path.join(folder, "Open Q88 Check.url"), "w", encoding="utf-8") as f:
            f.write(content)
    except OSError:
        pass


if __name__ == "__main__":
    # plain ASCII only here - some Windows console codepages can't encode
    # Korean and would crash the launcher before the server even starts
    print(f" * This PC: http://localhost:{PORT}/")
    print(f" * Colleagues on the same network: http://{_lan_ip()}:{PORT}/")
    print(" * Keep this window open - closing it stops the server for everyone.")
    _write_shortcut()
    threading.Thread(target=_open_browser, daemon=True).start()
    from waitress import serve
    serve(app, host="0.0.0.0", port=PORT)
