import copy
import re

from docx.table import _Row

SECTION_RE = re.compile(r"^\d+\.$")
FIELD_CODE_RE = re.compile(r"^\d+\.\d")


def unique_cells(row):
    """python-docx repeats a Cell wrapper for every grid column a merge spans;
    dedupe by the underlying <w:tc> element identity, preserving left-to-right order."""
    seen = []
    out = []
    for c in row.cells:
        if c._tc in seen:
            continue
        seen.append(c._tc)
        out.append(c)
    return out


def set_cell_text(cell, new_text):
    """Write new_text into a cell while preserving the formatting of its first
    run (font/size/color/bold), so saving only changes the value, not the layout."""
    paragraphs = cell.paragraphs
    first_p = paragraphs[0]

    # drop any extra paragraphs beyond the first
    for p in paragraphs[1:]:
        p._p.getparent().remove(p._p)

    runs = first_p.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r._r.getparent().remove(r._r)
    else:
        first_p.add_run(new_text)


def add_table_row(nested_table):
    """Append a new data row to a nested table by cloning the last data row's
    XML (so cell formatting matches exactly), then blanking its text."""
    rows = nested_table.rows
    if len(rows) < 2:
        return False
    template_row = rows[-1]
    new_tr = copy.deepcopy(template_row._tr)
    template_row._tr.addnext(new_tr)
    new_row = _Row(new_tr, nested_table)
    for cell in unique_cells(new_row):
        set_cell_text(cell, "")
    return True


def delete_table_row(nested_table, row_index):
    """Remove data row `row_index` (0-based, excluding the header row)."""
    rows = nested_table.rows
    target_idx = 1 + row_index
    if target_idx < 1 or target_idx >= len(rows):
        return False
    tr = rows[target_idx]._tr
    tr.getparent().remove(tr)
    return True


def _is_header_like(texts):
    for t in texts:
        t = t.strip()
        if not t or len(t) > 30 or any(ch.isdigit() for ch in t):
            return False
    return True


class Extraction:
    def __init__(self):
        self.display_rows = []  # ordered list of dicts for rendering
        self.cell_map = {}  # field id -> docx Cell object
        self.table_objs = {}  # table key -> docx Table object (only for editable/nested tables)
        self._next_id = 1
        self._next_key = 1

    def new_id(self):
        fid = f"f{self._next_id}"
        self._next_id += 1
        return fid

    def new_key(self):
        key = f"t{self._next_key}"
        self._next_key += 1
        return key

    def register_cell(self, column_header, cell):
        fid = self.new_id()
        self.cell_map[fid] = cell
        return {"id": fid, "column_header": column_header, "text": cell.text.strip()}

    def add_field(self, section, item_code, label, column_header, cell):
        fid = self.new_id()
        self.cell_map[fid] = cell
        self.display_rows.append({
            "type": "field",
            "id": fid,
            "section": section,
            "item_code": item_code,
            "label": label,
            "column_header": column_header,
            "text": cell.text.strip(),
        })
        return fid


def build_nested_table(ext, nested_table, section, item_code, label):
    """A real Word table embedded in a cell (mooring lines, pump data, ...).
    Row count varies per vessel, so these support add/delete row. Every
    column - including the first - is real, editable data with its own
    header (e.g. "Tank ID", "Nationality", "Pump Identity")."""
    if not nested_table.rows:
        return None
    header_cells = unique_cells(nested_table.rows[0])
    headers = [c.text.strip() for c in header_cells]

    key = ext.new_key()
    ext.table_objs[key] = nested_table
    table_rec = {
        "type": "table", "key": key, "section": section, "item_code": item_code,
        "label": label, "headers": headers, "editable_rows": True, "rows": [],
    }
    for row in nested_table.rows[1:]:
        cells = unique_cells(row)
        if not cells:
            continue
        row_label = cells[0].text.strip()
        row_cells = []
        for ci, c in enumerate(cells):
            header = headers[ci] if ci < len(headers) else ""
            row_cells.append(ext.register_cell(header, c))
        table_rec["rows"].append({"row_label": row_label, "cells": row_cells})
    return table_rec


def extract(doc):
    ext = Extraction()
    table = doc.tables[0]
    rows = table.rows

    rows_info = []
    for row in rows:
        cells = unique_cells(row)
        has_nested = any(len(c.tables) > 0 for c in cells)
        item_code_raw = cells[0].text.strip() if cells else ""
        rows_info.append({"cells": cells, "has_nested": has_nested, "item_code_raw": item_code_raw})

    current_section = None
    current_item_code = ""
    current_label = ""
    section_headers = None
    local_headers = None
    local_owner = None
    local_group_label = None

    current_table = None
    current_table_sig = None

    def flush_table():
        nonlocal current_table, current_table_sig
        if current_table is not None:
            ext.display_rows.append(current_table)
        current_table = None
        current_table_sig = None

    n = len(rows_info)
    for i, info in enumerate(rows_info):
        cells = info["cells"]

        if info["has_nested"]:
            flush_table()
            for c in cells:
                if len(c.tables) > 0:
                    rec = build_nested_table(ext, c.tables[0], current_section, current_item_code, current_label)
                    if rec:
                        ext.display_rows.append(rec)
            continue

        if len(cells) == 1:
            text = cells[0].text.strip()
            if text == "":
                continue
            flush_table()
            ext.display_rows.append({"type": "subheading", "text": text})
            continue

        item_code_raw = info["item_code_raw"]

        if SECTION_RE.match(item_code_raw):
            flush_table()
            label_text = cells[1].text.strip()
            current_section = label_text
            current_item_code = item_code_raw
            current_label = label_text
            extra = [c.text.strip() for c in cells[2:]]
            section_headers = extra if extra else None
            local_headers = None
            local_owner = None
            ext.display_rows.append({"type": "heading", "text": label_text})
            continue

        # pure spacer row (all cells blank) - not a real field, skip entirely
        if item_code_raw == "" and all(c.text.strip() == "" for c in cells):
            continue

        # normal field-ish row
        label_text = cells[1].text.strip()
        if item_code_raw and item_code_raw != current_item_code:
            current_item_code = item_code_raw
            local_headers = None
            local_owner = None
        current_label = label_text

        value_cells = cells[2:]

        if len(value_cells) == 0:
            flush_table()
            ext.display_rows.append({
                "type": "label", "item_code": current_item_code, "text": label_text,
            })
            continue

        # header-definer detection: >=2 short/header-like values, next non-empty
        # row continues the same item code (explicit repeat or blank continuation)
        if len(value_cells) >= 2 and _is_header_like([c.text for c in value_cells]):
            nxt = rows_info[i + 1] if i + 1 < n else None
            nxt_continues = bool(
                nxt and not nxt["has_nested"] and len(nxt["cells"]) >= 2
                and (nxt["item_code_raw"] == "" or nxt["item_code_raw"] == current_item_code)
            )
            if nxt_continues:
                flush_table()
                local_headers = [c.text.strip() for c in value_cells]
                local_owner = current_item_code
                local_group_label = label_text
                continue

        if local_owner == current_item_code and local_headers and len(value_cells) == len(local_headers):
            headers_to_use = local_headers
            sig = ("local", current_item_code)
            group_label = local_group_label
        elif section_headers and len(value_cells) == len(section_headers):
            headers_to_use = section_headers
            sig = ("section", current_section)
            group_label = current_section
        else:
            headers_to_use = None
            sig = None
            group_label = None

        if sig is not None:
            if current_table is None or current_table_sig != sig:
                flush_table()
                current_table = {
                    "type": "table", "key": ext.new_key(), "section": current_section,
                    "item_code": current_item_code if sig[0] == "local" else "",
                    "label": group_label, "headers": headers_to_use,
                    "editable_rows": False, "rows": [],
                }
                current_table_sig = sig
            row_cells = [ext.register_cell(headers_to_use[idx], vc) for idx, vc in enumerate(value_cells)]
            current_table["rows"].append({
                "row_label": label_text, "row_item_code": current_item_code, "cells": row_cells,
            })
            continue

        flush_table()

        # compound labels like "Date of last X/next Y due:" carry one segment
        # per value cell - split them so expiry keywords attach to the right value
        label_segments = None
        if len(value_cells) > 1:
            segments = [s.strip() for s in re.split(r"\s*/\s*", label_text)]
            if len(segments) == len(value_cells):
                label_segments = segments

        for idx, vc in enumerate(value_cells):
            col_header = None
            field_label = label_text
            if label_segments:
                field_label = label_segments[idx]
            elif len(value_cells) > 1:
                col_header = f"Value {idx + 1}"
            ext.add_field(current_section, current_item_code, field_label, col_header, vc)

    flush_table()
    return ext
