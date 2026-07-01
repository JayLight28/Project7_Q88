"""Synchronize value-cell font formatting (name/size/color/bold) across Q88
files that share the same official template, using one file as the reference.
Only visual formatting is touched - cell text/content is never modified here.
"""
from docx.oxml.ns import qn

from . import parser

RPR_FIELDS = ("ascii", "hAnsi", "cs", "sz", "szCs", "color", "bold")


def _rpr_of(paragraph):
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    return pPr.find(qn("w:rPr"))


def _read_rpr(rpr_el):
    if rpr_el is None:
        return {}
    spec = {}
    rFonts = rpr_el.find(qn("w:rFonts"))
    if rFonts is not None:
        for attr in ("ascii", "hAnsi", "cs"):
            val = rFonts.get(qn(f"w:{attr}"))
            if val:
                spec[attr] = val
    sz = rpr_el.find(qn("w:sz"))
    if sz is not None:
        spec["sz"] = sz.get(qn("w:val"))
    szCs = rpr_el.find(qn("w:szCs"))
    if szCs is not None:
        spec["szCs"] = szCs.get(qn("w:val"))
    color = rpr_el.find(qn("w:color"))
    if color is not None:
        spec["color"] = color.get(qn("w:val"))
    b = rpr_el.find(qn("w:b"))
    if b is not None:
        spec["bold"] = b.get(qn("w:val")) != "0"
    return spec


def capture_cell_font(cell):
    """Effective style for a cell: paragraph-mark defaults overlaid by the
    first run's explicit formatting (whichever pieces it overrides)."""
    if not cell.paragraphs:
        return {}
    para = cell.paragraphs[0]
    spec = _read_rpr(_rpr_of(para))
    if para.runs:
        run_rpr = para.runs[0]._r.find(qn("w:rPr"))
        spec.update(_read_rpr(run_rpr))
    return spec


def _write_rpr(rpr_el, spec):
    for child_tag in ("w:rFonts", "w:color", "w:sz", "w:szCs", "w:b"):
        existing = rpr_el.find(qn(child_tag))
        if existing is not None:
            rpr_el.remove(existing)

    if any(k in spec for k in ("ascii", "hAnsi", "cs")):
        rFonts = rpr_el.makeelement(qn("w:rFonts"), {})
        for attr in ("ascii", "hAnsi", "cs"):
            if attr in spec:
                rFonts.set(qn(f"w:{attr}"), spec[attr])
        rpr_el.append(rFonts)
    if "color" in spec:
        color = rpr_el.makeelement(qn("w:color"), {qn("w:val"): spec["color"]})
        rpr_el.append(color)
    if "sz" in spec:
        sz = rpr_el.makeelement(qn("w:sz"), {qn("w:val"): spec["sz"]})
        rpr_el.append(sz)
    if "szCs" in spec:
        szCs = rpr_el.makeelement(qn("w:szCs"), {qn("w:val"): spec["szCs"]})
        rpr_el.append(szCs)
    if spec.get("bold"):
        b = rpr_el.makeelement(qn("w:b"), {})
        rpr_el.append(b)


def apply_cell_font(cell, spec):
    if not spec or not cell.paragraphs:
        return
    para = cell.paragraphs[0]

    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = para._p.makeelement(qn("w:pPr"), {})
        para._p.insert(0, pPr)
    rpr = pPr.find(qn("w:rPr"))
    if rpr is None:
        rpr = pPr.makeelement(qn("w:rPr"), {})
        pPr.append(rpr)
    _write_rpr(rpr, spec)

    for run in para.runs:
        run_rpr = run._r.find(qn("w:rPr"))
        if run_rpr is None:
            run_rpr = run._r.makeelement(qn("w:rPr"), {})
            run._r.insert(0, run_rpr)
        _write_rpr(run_rpr, spec)


def _field_key(rec):
    return (rec["item_code"], rec["label"], rec["column_header"])


def _cell_keys(ext):
    """id -> style key, covering both flat field rows and cells nested inside
    table rows (certificates, mooring lines, engines, ...)."""
    keys = {}
    for r in ext.display_rows:
        if r["type"] == "field":
            keys[r["id"]] = _field_key(r)
        elif r["type"] == "table":
            for tr in r["rows"]:
                for c in tr["cells"]:
                    keys[c["id"]] = (r["key"], tr["row_label"], c["column_header"])
    return keys


def extract_style_map(doc):
    ext = parser.extract(doc)
    keys = _cell_keys(ext)
    style_map = {}
    for fid, cell in ext.cell_map.items():
        key = keys.get(fid)
        if key is None:
            continue
        style_map[key] = capture_cell_font(cell)
    return style_map


def apply_style_map(doc, style_map):
    ext = parser.extract(doc)
    keys = _cell_keys(ext)
    applied = 0
    for fid, cell in ext.cell_map.items():
        key = keys.get(fid)
        if key is None:
            continue
        spec = style_map.get(key)
        if spec:
            apply_cell_font(cell, spec)
            applied += 1
    return applied
