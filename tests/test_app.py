"""Route-level tests: tiny generated .docx in a tmp watch folder (selected via
the q88_folder cookie), exercising open/save/rename/lock-conflict/field_edit
and the Obsolete/ retention prune."""
import os
import sys
import time
from urllib.parse import quote

import docx as docxlib
import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import app as appmod
from q88 import locks, parser, state as statemod

FIXTURE_NAME = "Q88 V6 TEST 01 January 2026.docx"


def make_fixture(folder, name=FIXTURE_NAME):
    """Minimal Q88-shaped doc: one 3-column table of plain field rows (no
    section row, so the parser yields type='field' for every row).
    Field ids are assigned in row order: f1=1.1 date, f2=name, f3=flag, f4=expiry."""
    doc = docxlib.Document()
    table = doc.add_table(rows=0, cols=3)
    for code, label, value in [
        ("1.1", "Date updated:", "01-Jan-2026"),
        ("1.2", "Vessel's name:", "TEST SHIP"),
        ("1.5", "Flag:", "Panama"),
        ("1.9", "Certificate expiry", ""),
    ]:
        cells = table.add_row().cells
        cells[0].text = code
        cells[1].text = label
        cells[2].text = value
    path = os.path.join(str(folder), name)
    doc.save(path)
    return path


def set_cookie(c, key, value):
    try:
        c.set_cookie(key, value)  # Werkzeug >= 2.3
    except TypeError:
        c.set_cookie("localhost", key, value)


@pytest.fixture
def client(tmp_path):
    locks._locks.clear()
    appmod.app.config["TESTING"] = True
    with appmod.app.test_client() as c:
        set_cookie(c, "q88_folder", str(tmp_path))
        set_cookie(c, "q88_client_id", "tester-a")
        set_cookie(c, "q88_name", "Tester")
        yield c


def u(name):
    return quote(name)


def cell_text(path, fid):
    ext = parser.extract(docxlib.Document(path))
    return ext.cell_map[fid].text.strip()


def test_open_renders_and_acquires_lock(client, tmp_path):
    make_fixture(tmp_path)
    r = client.get("/open/" + u(FIXTURE_NAME))
    assert r.status_code == 200
    assert b"TEST SHIP" in r.data
    holder = locks.status(appmod._key(str(tmp_path), FIXTURE_NAME))
    assert holder and holder["name"] == "Tester"


def test_save_writes_cell_archives_and_records_history(client, tmp_path):
    path = make_fixture(tmp_path)
    r = client.post("/save/" + u(FIXTURE_NAME), data={"f_f3": "Greece"})
    assert r.status_code == 302
    assert cell_text(path, "f3") == "Greece"
    # pre-edit copy archived to Obsolete/
    assert os.path.exists(os.path.join(str(tmp_path), "Obsolete", FIXTURE_NAME))
    st = statemod.load_state(path)
    edits = [e for e in st["history"] if e["field_id"] == "f3"]
    assert edits and edits[-1]["old"] == "Panama" and edits[-1]["new"] == "Greece"
    assert edits[-1]["by"] == "Tester"


def test_save_renames_when_date_field_changes(client, tmp_path):
    path = make_fixture(tmp_path)
    r = client.post("/save/" + u(FIXTURE_NAME), data={"f_f1": "02-Feb-2026"})
    assert r.status_code == 302
    new_name = "Q88 V6 TEST 02 February 2026.docx"
    new_path = os.path.join(str(tmp_path), new_name)
    assert os.path.exists(new_path)
    assert not os.path.exists(path)
    assert new_name in r.headers["Location"] or u(new_name) in r.headers["Location"]
    # sidecar history moved with the file
    st = statemod.load_state(new_path)
    assert any(e["field_id"] == "__rename__" for e in st["history"])


def test_save_lost_lock_ajax_gets_409(client, tmp_path):
    make_fixture(tmp_path)
    locks.acquire(appmod._key(str(tmp_path), FIXTURE_NAME), "other-client", "Other")
    r = client.post(
        "/save/" + u(FIXTURE_NAME),
        data={"f_f3": "Greece"},
        headers={"X-Q88-Ajax": "1"},
    )
    assert r.status_code == 409
    assert r.get_json()["error"] == "locked"


def test_field_edit_updates_single_cell(client, tmp_path):
    path = make_fixture(tmp_path)
    r = client.post(
        "/field_edit/" + u(FIXTURE_NAME) + "/f4",
        data={"text": "01-Aug-2026", "na": "off"},
    )
    assert r.status_code == 200
    assert r.get_json()["ok"] is True
    assert cell_text(path, "f4") == "01-Aug-2026"


def test_heartbeat_reports_resumed_after_lock_loss(client, tmp_path):
    make_fixture(tmp_path)
    client.get("/open/" + u(FIXTURE_NAME))
    r = client.post("/heartbeat/" + u(FIXTURE_NAME))
    data = r.get_json()
    assert data["ok"] is True and data["resumed"] is False
    locks._locks.clear()  # simulate server restart / expiry sweep
    r2 = client.post("/heartbeat/" + u(FIXTURE_NAME))
    assert r2.get_json()["resumed"] is True


def test_prune_obsolete_deletes_old_keeps_newest_per_vessel(tmp_path):
    obs = tmp_path / "Obsolete"
    obs.mkdir()
    aaa_new = obs / "Q88 V6 AAA 01 January 2025.docx"
    aaa_old = obs / "Q88 V6 AAA 01 June 2023.docx"
    bbb_fresh = obs / "Q88 V6 BBB 01 July 2026.docx"
    ccc_lone = obs / "Q88 V6 CCC 01 January 2020.docx"
    for p in (aaa_new, aaa_old, bbb_fresh, ccc_lone):
        p.write_bytes(b"x")
    two_years = time.time() - 2 * 365 * 86400
    three_years = time.time() - 3 * 365 * 86400
    os.utime(aaa_new, (two_years, two_years))
    os.utime(aaa_old, (three_years, three_years))
    os.utime(ccc_lone, (three_years, three_years))

    appmod._prune_obsolete(str(obs))

    assert aaa_new.exists()  # newest AAA kept despite being past retention
    assert not aaa_old.exists()  # older AAA pruned
    assert bbb_fresh.exists()  # within retention
    assert ccc_lone.exists()  # only archive for CCC is never pruned


def test_prune_obsolete_groups_stamped_nonstandard_names(tmp_path):
    """Collision-stamped archives of a non-FILENAME_RE file must share one
    group - otherwise each stamp would be its own 'newest' and never age out."""
    obs = tmp_path / "Obsolete"
    obs.mkdir()
    newest = obs / "MyVessel.docx"
    stamped_old = obs / "MyVessel (20230101_120000).docx"
    stamped_older = obs / "MyVessel (20220601_090000).docx"
    for p in (newest, stamped_old, stamped_older):
        p.write_bytes(b"x")
    two_years = time.time() - 2 * 365 * 86400
    three_years = time.time() - 3 * 365 * 86400
    os.utime(stamped_old, (two_years, two_years))
    os.utime(stamped_older, (three_years, three_years))

    appmod._prune_obsolete(str(obs))

    assert newest.exists()
    assert not stamped_old.exists()
    assert not stamped_older.exists()
